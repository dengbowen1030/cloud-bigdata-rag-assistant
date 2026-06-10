"""Document loading entrypoint owned by member B.

This module converts PDF/DOCX/TXT course materials into the project-level
Document object plus raw text segments. It does not create embeddings or FAISS
indexes. The output is designed for ``rag.text_splitter`` and follows
``docs/module_contracts.md``.
"""

from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass, asdict
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

SUPPORTED_FILE_TYPES = {".pdf": "pdf", ".docx": "docx", ".txt": "txt"}


class DocumentProcessingError(RuntimeError):
    """Raised when parsing, reading, or validating an input document fails."""


@dataclass
class TextSegment:
    """A parsed text block with source location metadata."""

    text: str
    page: Optional[int]
    section: Optional[str] = None
    segment_index: int = 1

    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


def detect_file_type(file_path: str | Path) -> str:
    """Return normalized file type: pdf/docx/txt.

    Raises:
        DocumentProcessingError: when the extension is unsupported.
    """

    suffix = Path(file_path).suffix.lower()
    if suffix not in SUPPORTED_FILE_TYPES:
        raise DocumentProcessingError(
            f"Unsupported file type: {suffix or '<no extension>'}. "
            "Supported types are PDF, DOCX and TXT."
        )
    return SUPPORTED_FILE_TYPES[suffix]


def build_document_id(file_path: str | Path, prefix: str = "doc") -> str:
    """Build a stable document id from filename, size and modified time."""

    path = Path(file_path)
    stat = path.stat()
    raw = f"{path.name}:{stat.st_size}:{int(stat.st_mtime)}"
    digest = hashlib.md5(raw.encode("utf-8")).hexdigest()[:10]
    safe_stem = re.sub(r"[^a-zA-Z0-9_]+", "_", path.stem).strip("_").lower()
    if not safe_stem:
        safe_stem = "file"
    return f"{prefix}_{safe_stem}_{digest}"


def build_document(file_path: str | Path, document_id: Optional[str] = None) -> Dict[str, Any]:
    """Create a Document dict with chunk_count temporarily set to 0."""

    path = Path(file_path)
    if not path.exists() or not path.is_file():
        raise DocumentProcessingError(f"File not found: {path}")
    if path.stat().st_size <= 0:
        raise DocumentProcessingError(f"Empty file: {path.name}")

    file_type = detect_file_type(path)
    return {
        "document_id": document_id or build_document_id(path),
        "filename": path.name,
        "file_type": file_type,
        "file_size": path.stat().st_size,
        "status": "processed",
        "chunk_count": 0,
        "created_at": datetime.now().replace(microsecond=0).isoformat(),
    }


def _extract_pdf_segments(file_path: Path) -> List[TextSegment]:
    """Extract PDF text page by page using pypdf, falling back to PyPDF2."""

    reader_cls = None
    try:
        from pypdf import PdfReader  # type: ignore

        reader_cls = PdfReader
    except ImportError:
        try:
            from PyPDF2 import PdfReader  # type: ignore

            reader_cls = PdfReader
        except ImportError as exc:
            raise DocumentProcessingError(
                "PDF parsing requires dependency 'pypdf' or 'PyPDF2'. "
                "Run: pip install pypdf"
            ) from exc

    try:
        reader = reader_cls(str(file_path))
        segments: List[TextSegment] = []
        for idx, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                segments.append(TextSegment(text=text, page=idx, segment_index=idx))
        return segments
    except Exception as exc:  # pragma: no cover - depends on PDF parser internals
        raise DocumentProcessingError(f"Failed to parse PDF {file_path.name}: {exc}") from exc


def _extract_docx_segments(file_path: Path) -> List[TextSegment]:
    """Extract DOCX text by paragraph and table cell."""

    try:
        from docx import Document as DocxDocument  # type: ignore
    except ImportError as exc:
        raise DocumentProcessingError(
            "DOCX parsing requires dependency 'python-docx'. "
            "Run: pip install python-docx"
        ) from exc

    try:
        doc = DocxDocument(str(file_path))
        segments: List[TextSegment] = []
        seg_index = 1
        current_section: Optional[str] = None

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = getattr(paragraph.style, "name", "") or ""
            if style_name.lower().startswith("heading") or re.match(r"^第[一二三四五六七八九十0-9]+[章节篇部分]", text):
                current_section = text[:80]
            segments.append(
                TextSegment(
                    text=text,
                    page=None,
                    section=current_section,
                    segment_index=seg_index,
                )
            )
            seg_index += 1

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    segments.append(
                        TextSegment(
                            text=" | ".join(cells),
                            page=None,
                            section=current_section,
                            segment_index=seg_index,
                        )
                    )
                    seg_index += 1
        return segments
    except Exception as exc:  # pragma: no cover - depends on python-docx internals
        raise DocumentProcessingError(f"Failed to parse DOCX {file_path.name}: {exc}") from exc


def _read_txt_with_fallback(file_path: Path) -> str:
    encodings = ["utf-8", "utf-8-sig", "gb18030", "gbk", "big5"]
    for encoding in encodings:
        try:
            return file_path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentProcessingError(f"Failed to decode TXT file: {file_path.name}")


def _extract_txt_segments(file_path: Path) -> List[TextSegment]:
    """Extract TXT by paragraphs while preserving paragraph boundaries."""

    text = _read_txt_with_fallback(file_path)
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    if not paragraphs and text.strip():
        paragraphs = [text.strip()]
    return [TextSegment(text=p, page=None, segment_index=i) for i, p in enumerate(paragraphs, start=1)]


def load_document(file_path: str | Path, document_id: Optional[str] = None) -> Tuple[Dict[str, Any], str, List[Dict[str, Any]]]:
    """Load a document and return ``(document, raw_text, segments)``.

    Args:
        file_path: Path under ``uploads/raw/`` or ``data/raw/``.
        document_id: Optional external id from backend A. If omitted, a stable id
            is generated from file metadata.

    Returns:
        document: Document dict matching ``docs/module_contracts.md``.
        raw_text: Joined extracted text.
        segments: List of dicts with text/page/section metadata for chunking.

    Raises:
        DocumentProcessingError: when parsing fails or no text is extracted.
    """

    path = Path(file_path)
    document = build_document(path, document_id=document_id)
    file_type = document["file_type"]

    if file_type == "pdf":
        parsed_segments = _extract_pdf_segments(path)
    elif file_type == "docx":
        parsed_segments = _extract_docx_segments(path)
    elif file_type == "txt":
        parsed_segments = _extract_txt_segments(path)
    else:  # Defensive guard; detect_file_type should already reject this.
        raise DocumentProcessingError(f"Unsupported file type: {file_type}")

    parsed_segments = [seg for seg in parsed_segments if seg.text and seg.text.strip()]
    if not parsed_segments:
        raise DocumentProcessingError(f"No extractable text found in {path.name}")

    raw_text = "\n\n".join(seg.text for seg in parsed_segments)
    return document, raw_text, [seg.to_dict() for seg in parsed_segments]
